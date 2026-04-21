# Importaciones relativas del proyecto
from ..models import Universidad, Tema, Curso, Pregunta, UserProfile
from ..forms import Pregunta, PreguntaForm
import uuid

# Django - shortcuts y decoradores
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.admin.views.decorators import staff_member_required
from django.views.decorators.http import require_POST
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.core.paginator import Paginator
from django.http import HttpResponse, FileResponse, Http404, JsonResponse
from django.conf import settings
from django.utils import timezone
from django.utils.text import slugify
from django.views.decorators.clickjacking import xframe_options_exempt
from django.contrib import messages
from django.urls import reverse
from django.core.files.base import ContentFile
from django.views.decorators.csrf import csrf_exempt
from django.core.exceptions import PermissionDenied

# Python estándar
import os, io, requests
import logging
from collections import defaultdict
from datetime import timedelta

# Librerías de terceros para manejo de documentos DOCX
from docx import Document
from docxcompose.composer import Composer
from docx.enum.text import WD_COLOR_INDEX

try:
    from docxcompose.composer import ImportFormatMode
except ImportError:
    ImportFormatMode = None

from docx.shared import Pt, Inches, Cm  # Tamaño de fuente y márgenes
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn

# Importación de vistas de autenticación propias
from .auth_views import exclude_supervisor, role_required

import json
import time
import jwt

# Gestión de Preguntas
from django.core.paginator import Paginator

@login_required
@role_required('admin', 'user')
def pregunta_list(request):
    if request.user.is_superuser:
        qs_base = Pregunta.objects.all()
    else:
        user_profile = get_object_or_404(UserProfile, user=request.user)
        qs_base = Pregunta.objects.filter(usuario=user_profile)

    tiempo_filtro = request.GET.get('tiempo_filtro')
    
    if tiempo_filtro:
        try:
            minutos = int(tiempo_filtro)
            limite_tiempo = timezone.now() - timedelta(minutes=minutos)
            qs_base = qs_base.filter(fecha_creacion__gte=limite_tiempo)
        except ValueError:
            pass 
    elif not request.user.is_superuser:
        limite_default = timezone.now() - timedelta(days=1)
        qs_base = qs_base.filter(fecha_creacion__gte=limite_default)

    universidades_qs = Universidad.objects.filter(pregunta__in=qs_base).distinct()

    universidad_id = request.GET.get('universidad')
    cursos_para_uni = []
    if universidad_id:
        cursos_para_uni = Curso.objects.filter(
            pregunta__in=qs_base, 
            pregunta__universidad_id=universidad_id
        ).distinct()

    curso_id = request.GET.get('curso')
    temas_para_curso = []
    if curso_id:
        temas_qs = Tema.objects.filter(curso_id=curso_id, pregunta__in=qs_base)
        if universidad_id:
            temas_qs = temas_qs.filter(pregunta__universidad_id=universidad_id)
        
        temas_para_curso = temas_qs.distinct()
        
    tema_id = request.GET.get('tema')
    nivel = request.GET.get('nivel')

    qs = qs_base

    if universidad_id:
        qs = qs.filter(universidad_id=universidad_id)
    if curso_id:
        qs = qs.filter(curso_id=curso_id)
    if tema_id:
        qs = qs.filter(tema_id=tema_id)
    if nivel:
        qs = qs.filter(nivel=nivel)

    qs = qs.order_by('-fecha_creacion')

    paginator = Paginator(qs, 100)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    context = {
        'page_obj': page_obj,
        'universidades': universidades_qs,
        'cursos_para_uni': cursos_para_uni,
        'temas_para_curso': temas_para_curso,
        'universidad_filter': universidad_id,
        'curso_filter': curso_id,
        'tema_filter': tema_id,
        'nivel_filter': nivel,
        'tiempo_filtro': tiempo_filtro,
        'total_preguntas': qs.count(), 
    }
    return render(request, 'Preguntas/preguntas/pregunta_list.html', context)

@login_required
@role_required('admin', 'supervisor')
def pregunta_list_supervisor(request):
    qs = Pregunta.objects.select_related('usuario__user', 'universidad', 'curso', 'tema').all()

    buscar_nombre = request.GET.get('nombre')
    if buscar_nombre:
        qs = qs.filter(nombre__icontains=buscar_nombre)

    buscar_usuario = request.GET.get('usuario')
    if buscar_usuario:
        qs = qs.filter(usuario_id=buscar_usuario)

    qs = qs.order_by('-fecha_creacion')

    paginator = Paginator(qs, 100)
    page = request.GET.get('page')
    qs_paginated = paginator.get_page(page)

    from django.contrib.auth.models import User
    usuarios = User.objects.all()

    context = {
        'preguntas': qs_paginated,
        'buscar_nombre': buscar_nombre,
        'usuarios': usuarios,
        'buscar_usuario': buscar_usuario,
    }

    return render(request, 'Preguntas/preguntas/lista_supervisor.html', context)

@login_required
@exclude_supervisor
@role_required('admin', 'user')
def pregunta_create(request):
    if request.method == 'POST':
        form = PreguntaForm(request.POST, request.FILES)
        if form.is_valid():
            pregunta = form.save(commit=False)
            user_profile = UserProfile.objects.get(user=request.user)
            pregunta.usuario = user_profile

            count = Pregunta.objects.filter(
                universidad=pregunta.universidad,
                curso=pregunta.curso,
                tema=pregunta.tema,
                nivel=pregunta.nivel
            ).count() + 1

            pregunta.nombre = f"{pregunta.universidad.id}{pregunta.curso.id}{pregunta.tema.id}{pregunta.nivel}{count}"
            pregunta.save()

            if 'contenido' in request.FILES:
                pregunta.contenido = request.FILES['contenido']
                pregunta.save()

            messages.success(request, 'Pregunta creada exitosamente.')

            data = request.POST.copy()
            data['nombre'] = ''
            nuevo_formulario = PreguntaForm(data, is_update=False)
            nuevo_formulario.fields['contenido'].required = False

            return render(request, 'Preguntas/preguntas/pregunta_form.html', {
                'form': nuevo_formulario,
                'title': 'Nueva Pregunta'
            })

        else:
            return render(request, 'Preguntas/preguntas /pregunta_form.html', {
                'form': form,
                'title': 'Nueva Pregunta'
            })
    else:
        form = PreguntaForm()

    return render(request, 'Preguntas/pregunta_form.html', {
        'form': form,
        'title': 'Nueva Pregunta'
    })

@login_required
@exclude_supervisor
@role_required('admin', 'user')
def pregunta_update(request, pk):
    try:
        user_profile = UserProfile.objects.get(user=request.user)
    except UserProfile.DoesNotExist:
        messages.error(request, 'No se encontró el perfil de usuario.')
        return redirect('pregunta-list')

    if request.user.is_superuser:
        pregunta = get_object_or_404(Pregunta, pk=pk)
    else:
        pregunta = get_object_or_404(Pregunta, pk=pk, usuario=user_profile)
        
    if request.method == 'POST':
        form = PreguntaForm(request.POST, request.FILES, instance=pregunta, is_update=True)
        
        if form.is_valid():
            nuevo_archivo = request.FILES.get('contenido')

            # Eliminar PDF existente si hay cambios
            pdf_dir = os.path.join(settings.MEDIA_ROOT, 'pdfs')
            safe_filename = sanitize_filename(pregunta.nombre)
            pdf_filename = f"{safe_filename}_{pregunta.id}.pdf"
            pdf_path = os.path.join(pdf_dir, pdf_filename)
            
            if os.path.exists(pdf_path):
                try:
                    os.remove(pdf_path)
                    logger.info(f"PDF eliminado por actualización de pregunta: {pdf_path}")
                except Exception as e:
                    logger.error(f"No se pudo eliminar PDF: {e}")

            # Actualizar campos editables
            pregunta.respuesta = form.cleaned_data['respuesta']
            pregunta.nivel = form.cleaned_data['nivel']
            pregunta.tiene_solucion = form.cleaned_data['tiene_solucion']
            
            if nuevo_archivo:
                pregunta.contenido = nuevo_archivo
                logger.info(f"📝 Archivo actualizado para pregunta {pk}: {nuevo_archivo.name}")
            
            pregunta.save()

            messages.success(request, 'Pregunta actualizada con éxito.')
            return redirect('pregunta-list')
        else:
            # Mostrar errores del formulario
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f"Error en {field}: {error}")
    else:
        form = PreguntaForm(instance=pregunta, is_update=True)

    return render(request, 'Preguntas/preguntas/pregunta_form.html', {
        'form': form,
        'pregunta': pregunta,
        'title': 'Editar Pregunta',
        'is_update': True,
        'current_file': pregunta.contenido.name if pregunta.contenido else None
    })

@require_POST
@login_required
@role_required('admin')
def actualizar_rapido_pregunta(request):
    try:
        pregunta_id = request.POST.get('id')
        tiene_solucion = request.POST.get('tiene_solucion', 'false') == 'true'
        alternativa = request.POST.get('alternativa', '').upper()

        logger.info(f"Actualización rápida solicitada - User: {request.user}, Pregunta ID: {pregunta_id}")

        if not pregunta_id:
            logger.warning("ID de pregunta no proporcionado")
            return JsonResponse({'success': False, 'error': 'ID de pregunta no proporcionado'}, status=400)

        if alternativa not in ['A', 'B', 'C', 'D', 'E']:
            logger.warning(f"Alternativa inválida recibida: {alternativa}")
            return JsonResponse({'success': False, 'error': 'Alternativa inválida'}, status=400)

        try:
            pregunta = Pregunta.objects.select_related('usuario__user').get(id=pregunta_id)
            
            # Verificar permisos mejorado
            if not (request.user.is_superuser or 
                   (hasattr(pregunta, 'usuario') and 
                    pregunta.usuario and 
                    pregunta.usuario.user == request.user)):
                logger.warning(f"Intento de edición no autorizado. User: {request.user}, Pregunta: {pregunta_id}")
                return JsonResponse(
                    {'success': False, 'error': 'No tienes permisos para editar esta pregunta'},
                    status=403
                )

            # Actualizar campos
            pregunta.tiene_solucion = tiene_solucion
            pregunta.respuesta = alternativa
            pregunta.save(update_fields=['tiene_solucion', 'respuesta'])

            logger.info(f"Pregunta {pregunta_id} actualizada correctamente por {request.user}")
            return JsonResponse({'success': True})
            
        except Pregunta.DoesNotExist:
            logger.error(f"Pregunta no encontrada: {pregunta_id}")
            return JsonResponse({'success': False, 'error': 'Pregunta no encontrada'}, status=404)
            
    except Exception as e:
        logger.error(f"Error en actualización rápida: {str(e)}", exc_info=True)
        return JsonResponse({'success': False, 'error': 'Error interno del servidor'}, status=500)
    
@login_required
@exclude_supervisor
@role_required('admin', 'user')
def pregunta_delete(request, pk):
    try:
        user_profile = UserProfile.objects.get(user=request.user)
    except UserProfile.DoesNotExist:
        messages.error(request, 'No se encontró el perfil de usuario.')
        return redirect('pregunta-list')

    if request.user.is_superuser:
        pregunta = get_object_or_404(Pregunta, pk=pk)
    else:
        pregunta = get_object_or_404(Pregunta, pk=pk, usuario=user_profile)

    # Capturamos la URL anterior solo para el GET (para mostrar "volver")
    referer = request.META.get('HTTP_REFERER', reverse('pregunta-list'))

    if request.method == 'POST':
        pregunta.delete()
        messages.success(request, 'Pregunta eliminada exitosamente.')
        return redirect('pregunta-list')  # Redirige siempre a la lista

    return render(request, 'Preguntas/preguntas/pregunta_confirm_delete.html', {
        'pregunta': pregunta,
        'volver_url': referer, 
    })

@login_required
@exclude_supervisor
@role_required('admin', 'user')
def eliminar_preguntas(request):
    if request.method == 'POST':
        pregunta_ids = request.POST.getlist('preguntas')
        
        try:
            user_profile = UserProfile.objects.get(user=request.user)
        except UserProfile.DoesNotExist:
            messages.error(request, "No se encontró el perfil de usuario.")
            return redirect('pregunta-list')

        # Verificar que todas las preguntas pertenecen al usuario (a menos que sea superuser)
        if request.user.is_superuser:
            preguntas = Pregunta.objects.filter(id__in=pregunta_ids)
        else:
            preguntas = Pregunta.objects.filter(id__in=pregunta_ids, usuario=user_profile)

        count = preguntas.count()
        if count == 0:
            messages.error(request, 'No se encontraron preguntas para eliminar.')
            return redirect('pregunta-list')

        # Eliminar las preguntas
        preguntas.delete()
        messages.success(request, f'Se eliminaron {count} pregunta(s) correctamente.')
        
        return redirect('pregunta-list')

    # Si no es POST, redirigir
    return redirect('pregunta-list')

#desde aquí empecé a modificar lo del formato de las preguntas
#para darle 2 columnas al doc final
def set_tres_columns(section):
    sectPr = section._sectPr  # Obtener el elemento de la sección
    cols = OxmlElement('w:cols')
    cols.set(ns.qn('w:num'), '3')  # Establecer dos columnas
    sectPr.append(cols)

def set_margenes(section):
    """Configura los márgenes del documento según lo solicitado."""
    section.top_margin = Inches(2 / 2.54)  # 2 cm
    section.left_margin = Inches(0.76 / 2.54)  # 0.76 cm
    section.right_margin = Inches(0.76 / 2.54)  # 0.76 cm
    section.bottom_margin = Inches(3.25 / 2.54)  # 3.25 cm

def aplicar_formato_texto(doc):
    """Aplica Arial Narrow y tamaño 9 pt a todo el contenido del documento."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial Narrow"
            run.font.size = Pt(9)
            r = run._element
            r.rPr.rFonts.set(qn("w:eastAsia"), "Arial Narrow")
    for style in doc.styles:
        if style.type == 1:  # Solo afecta estilos de párrafo
            if style.name.lower() in ["list paragraph", "lista numerada", "lista con viñetas"]:
                style.font.name = "Arial Narrow"
                style.font.size = Pt(9)

def eliminar_ultimo_parrafo_si_vacio(document):
    if len(document.paragraphs) == 0:
        return

    last_paragraph = document.paragraphs[-1]
    if not last_paragraph.text.strip():
        p_element = last_paragraph._element
        p_element.getparent().remove(p_element)

def combinar_documentos(preguntas):
    """Combina documentos mostrando solo nombres de preguntas sin espacios adicionales"""
    master_doc = Document()
    composer = Composer(master_doc)
    
    # Configurar documento con 3 columnas y márgenes
    set_tres_columns(master_doc.sections[0])
    set_margenes(master_doc.sections[0])
    
    # Eliminar espaciado por defecto en el documento base
    style = master_doc.styles['Normal']
    font = style.font
    font.name = 'Arial Narrow'
    font.size = Pt(9)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.0  # Espaciado simple

    # Agrupar preguntas por curso y tema (solo para orden)
    preguntas_ordenadas = defaultdict(lambda: defaultdict(list))
    for pregunta in preguntas:
        if pregunta.contenido and hasattr(pregunta.contenido, 'path'):
            preguntas_ordenadas[pregunta.curso.nombre][pregunta.tema.nombre].append(pregunta)

    # Procesar preguntas
    for curso, temas in sorted(preguntas_ordenadas.items()):
        for tema, preguntas_tema in sorted(temas.items()):
            for pregunta in preguntas_tema:
                try:
                    # Agregar nombre de pregunta sin espaciado
                    p = master_doc.add_paragraph(style='Normal')
                    run = p.add_run(f"Pregunta: {pregunta.nombre}")
                    run.bold = True
                    
                    # Procesar contenido de la pregunta
                    sub_doc = Document(pregunta.contenido.path)
                    
                    # Eliminar configuraciones de sección y ajustar formato
                    for element in sub_doc.element.body:
                        if element.tag.endswith('sectPr'):
                            sub_doc.element.body.remove(element)
                    
                    # Aplicar formato sin espacios a todos los párrafos
                    for para in sub_doc.paragraphs:
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(0)
                        for run in para.runs:
                            run.font.name = "Arial Narrow"
                            run.font.size = Pt(9)
                    
                    # Combinar documentos
                    if ImportFormatMode is not None:
                        composer.append(sub_doc, import_format=ImportFormatMode.KEEP_SOURCE_FORMATTING)
                    else:
                        composer.append(sub_doc)
                        
                except Exception as e:
                    logger.error(f"Error procesando pregunta {pregunta.id}: {str(e)}")
                    continue

    # Eliminar posibles párrafos vacíos finales
    eliminar_ultimo_parrafo_si_vacio(composer.doc)

    buffer = io.BytesIO()
    composer.save(buffer)
    buffer.seek(0)
    return buffer

logger = logging.getLogger(__name__)

def sanitize_filename(filename):
    filename = slugify(filename, allow_unicode=False)
    if len(filename) > 100:
        filename = filename[:100]
    return filename or "documento"

@xframe_options_exempt
@login_required
def vista_previa(request, pk):
    try:
        pregunta = Pregunta.objects.get(pk=pk)
    except Pregunta.DoesNotExist:
        logger.warning(f"Pregunta {pk} no encontrada")
        raise Http404("La pregunta no existe")

    # Validar que el archivo DOCX existe
    docx_path = os.path.join(settings.MEDIA_ROOT, str(pregunta.contenido))
    if not os.path.exists(docx_path):
        logger.error(f"Archivo DOCX no encontrado: {docx_path}")
        raise Http404("El archivo DOCX no existe")

    # Crear directorio para PDFs si no existe
    pdf_dir = os.path.join(settings.MEDIA_ROOT, 'pdfs')
    os.makedirs(pdf_dir, exist_ok=True)

    # Generar nombre único para el PDF basado en la pregunta
    safe_filename = sanitize_filename(pregunta.nombre)
    pdf_filename = f"{safe_filename}_{pregunta.id}.pdf"
    pdf_path = os.path.join(pdf_dir, pdf_filename)

    # CONVERSIÓN SOLO SI EL PDF NO EXISTE
    if not os.path.exists(pdf_path):
        logger.info(f"Iniciando conversión: {docx_path} -> {pdf_path}")
        
        try:
            # Conversión con Aspose Words
            import aspose.words as aw
            
            # Cargar documento DOCX
            doc = aw.Document(docx_path)
            
            # Configurar opciones de guardado
            save_options = aw.saving.PdfSaveOptions()
            save_options.compliance = aw.saving.PdfCompliance.PDF17
            save_options.preserve_form_fields = True
            save_options.jpeg_quality = 90
            
            # Convertir a PDF
            doc.save(pdf_path, save_options)
            
            logger.info(f"✅ Conversión exitosa para pregunta {pk}")
            
        except ImportError:
            logger.error("❌ Aspose Words no está instalado")
            raise Http404("Error: Aspose Words no está disponible en el sistema")
        except Exception as e:
            logger.error(f"❌ Error en conversión: {e}")
            # Limpiar archivo parcial si existe
            if os.path.exists(pdf_path):
                try:
                    os.remove(pdf_path)
                except:
                    pass
            raise Http404(f"Error al convertir archivo: {str(e)}")
    else:
        logger.info(f"✅ Usando PDF existente para pregunta {pk}")

    # VALIDACIONES FINALES
    if not os.path.exists(pdf_path):
        logger.error(f"❌ PDF no se generó correctamente: {pdf_path}")
        raise Http404("No se pudo generar el archivo PDF")

    try:
        file_size = os.path.getsize(pdf_path)
        if file_size == 0:
            logger.error(f"❌ PDF generado está vacío: {pdf_path}")
            os.remove(pdf_path)
            raise Http404("El archivo PDF generado está vacío")
        elif file_size < 100:
            logger.warning(f"⚠️ PDF sospechosamente pequeño ({file_size} bytes): {pdf_path}")
        
        logger.info(f"📄 Sirviendo PDF: {pdf_path} ({file_size} bytes)")
        
    except OSError as e:
        logger.error(f"❌ Error verificando PDF: {e}")
        raise Http404("Error al verificar el archivo PDF")

    return serve_pdf_file(pdf_path, pregunta.nombre)

def serve_pdf_file(pdf_path, original_name):
    try:
        pdf_file = open(pdf_path, 'rb')
        safe_download_name = sanitize_filename(original_name)
        
        # Obtener timestamp de modificación del archivo
        last_modified = int(os.path.getmtime(pdf_path))
        download_filename = f"{safe_download_name}_{last_modified}.pdf"
        
        response = FileResponse(
            pdf_file,
            content_type='application/pdf',
            filename=download_filename
        )
        
        # Headers para evitar cache
        response['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
        response['Pragma'] = 'no-cache'
        response['Expires'] = '0'
        response['X-Accel-Expires'] = '0'
        
        return response
    except Exception as e:
        logger.error(f"Error al servir PDF: {e}")
        raise Http404("Error al acceder al archivo PDF")

def cleanup_old_pdfs():
    """Limpia PDFs antiguos para ahorrar espacio"""
    try:
        pdf_dir = os.path.join(settings.MEDIA_ROOT, 'pdfs')
        if not os.path.exists(pdf_dir):
            return
        
        import time
        current_time = time.time()
        week_ago = current_time - (3 * 24 * 60 * 60)  # 7 días
        
        for filename in os.listdir(pdf_dir):
            file_path = os.path.join(pdf_dir, filename)
            if os.path.isfile(file_path):
                file_mtime = os.path.getmtime(file_path)
                if file_mtime < week_ago:
                    try:
                        os.remove(file_path)
                        logger.info(f"PDF antiguo eliminado: {filename}")
                    except Exception as e:
                        logger.warning(f"No se pudo eliminar {filename}: {e}")
                        
    except Exception as e:
        logger.error(f"Error en limpieza de PDFs: {e}")

@staff_member_required
@exclude_supervisor
@login_required
@role_required('admin')
def todas_las_preguntas(request):
    qs = Pregunta.objects.all()

    # Leer filtros del GET
    buscar_nombre = request.GET.get('nombre')
    usuario_id    = request.GET.get('usuario')

    # Aplicar filtros si existen
    if buscar_nombre:
        qs = qs.filter(nombre__icontains=buscar_nombre)
    if usuario_id:
        qs = qs.filter(usuario_id=usuario_id)

    # Ordenar por fecha reciente
    qs = qs.order_by('-fecha_creacion')

    # Paginación
    paginator = Paginator(qs, 20)
    page = request.GET.get('page')
    qs_paginated = paginator.get_page(page)

    context = {
        'total_preguntas': qs.count(),
        'preguntas': qs_paginated,
        'usuarios': User.objects.all(),
        'buscar_nombre': buscar_nombre,
        'usuario_id': usuario_id,
        'modo_admin': True,
        'now': timezone.now(),

    }

    return render(request, 'Preguntas/admin/todas_las_preguntas.html', context)

@login_required
@exclude_supervisor
def descargar_preguntas(request):
    pregunta_ids = request.POST.getlist('preguntas')
    
    try:
        user_profile = UserProfile.objects.get(user=request.user)
    except UserProfile.DoesNotExist:
        messages.error(request, "No se encontró el perfil de usuario.")
        return redirect('pregunta-list')

    preguntas = Pregunta.objects.filter(id__in=pregunta_ids, usuario=user_profile)
    
    if not preguntas:
        messages.error(request, 'No se encontraron preguntas para descargar.')
        return redirect('pregunta-list')
    
    buffer = combinar_documentos(preguntas)
    
    response = HttpResponse(
        buffer,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename="preguntas_combinadas.docx"'
    
    return response

## EDITOR EN LINEA CON ONLYFFICE :) xd
def crear_docx_minimo(ancho_cm=7):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(ancho_cm)
    section.page_height = Cm(29.7) 
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Cm(0.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial Narrow'
    font.size = Pt(9)
    
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), 'Arial Narrow')
    rFonts.set(qn('w:hAnsi'), 'Arial Narrow')

    doc.add_paragraph("")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

def generar_token_office(request, pregunta, modo='edit', es_solucion=False):
    dominio_publico = settings.SITE_DOMAIN

    ip_interna = "http://192.168.18.20:8003"
    
    archivo = pregunta.solucion_archivo if es_solucion else pregunta.contenido
    if not archivo:
        logger.error(f"Error: Intento de editar archivo inexistente para pregunta {pregunta.id}")
        return None, None

    file_url = f"{ip_interna}{archivo.url}"
    tipo_str = 'sol' if es_solucion else 'pre'
    
    try:
        version_key = int(os.path.getmtime(archivo.path))
    except Exception as e:
        logger.warning(f"No se pudo obtener mtime: {e}. Usando ID como fallback.")
        version_key = pregunta.id

    # La KEY DEBE ser distinta para PRE y SOL
    document_key = f"{tipo_str.upper()}_{pregunta.id}_{uuid.uuid4().hex}"
    
    # Callback dinámico
    callback_url = f"{ip_interna}/onlyoffice/callback/?id={pregunta.id}&tipo={tipo_str}"
    
    logger.info(f"--- GENERANDO TOKEN ONLYOFFICE ---")
    logger.info(f"Modo: {modo} | Tipo: {tipo_str}")
    logger.info(f"URL Archivo: {file_url}")
    logger.info(f"URL Callback: {callback_url}")
    logger.info(f"Document Key: {document_key}")
    
    payload = {
        "iat": int(time.time()),
        "exp": int(time.time()) + 3600,
        "document": {
            "fileType": "docx",
            "key": document_key, 
            "title": f"{tipo_str.upper()}_{pregunta.nombre}.docx",
            "url": file_url,
            "permissions": {
                "edit": modo == 'edit',
                "download": True,
                "copy": True,
            }
        },
        "editorConfig": {
            "mode": modo,
            "lang": "es",
            "callbackUrl": callback_url if modo == 'edit' else None,
            "customization": {
                "forcesave": True,
                "autosave": True,
                "compactHeader": True,
                "toolbar": modo == 'edit',
            }
        }
    }
    token = jwt.encode(payload, settings.ONLYOFFICE_JWT_SECRET, algorithm="HS256")
    
    payload_js = {
        "document": json.dumps(payload["document"]),
        "editorConfig": json.dumps(payload["editorConfig"])
    }

    return token, payload_js

# --- VISTAS DEL FLUJO ---

@login_required
@role_required('admin', 'user')
def configurar_contexto(request):
    """Pantalla inicial para definir Universidad, Curso, etc."""
    if request.method == "POST":
        request.session['admin_uni'] = request.POST.get('universidad')
        request.session['admin_curso'] = request.POST.get('curso')
        request.session['admin_tema'] = request.POST.get('tema')
        request.session['admin_nivel'] = request.POST.get('nivel')
        return redirect('flujo_carga_continua')
    
    return render(request, 'Preguntas/preguntas/configurar_contexto.html', {
        'universidades': Universidad.objects.all().order_by('nombre'),
        'cursos': Curso.objects.all().order_by('nombre'),
        'temas': Tema.objects.all().order_by('nombre'),
    })

@login_required
@role_required('admin', 'user')
def flujo_carga_continua(request):
    """Crea la pregunta y abre el editor de 7cm."""
    uni_id = request.session.get('admin_uni')
    if not uni_id: 
        return redirect('configurar_contexto')

    nueva_pregunta = Pregunta.objects.create(
        universidad_id=uni_id,
        curso_id=request.session.get('admin_curso'),
        tema_id=request.session.get('admin_tema'),
        nivel=request.session.get('admin_nivel'),
        usuario=request.user.userprofile
    )
    
    # Guardar archivo inicial
    nueva_pregunta.contenido.save(
        f"{nueva_pregunta.nombre}.docx", 
        ContentFile(crear_docx_minimo(7))
    )

    token, payload = generar_token_office(request, nueva_pregunta, modo='edit')

    return render(request, 'Preguntas/preguntas/editor_continuo.html', {
        'pregunta': nueva_pregunta,
        'DOC_TOKEN': token,
        'PAYLOAD': payload,
        'ONLYOFFICE_API_URL': settings.ONLYOFFICE_API_URL
    })

@csrf_exempt
def onlyoffice_callback(request):
    """Guarda el archivo y detecta la clave resaltada."""
    if request.method != "POST":
        return JsonResponse({"error": 1, "message": "Método no permitido"})

    try:
        data = json.loads(request.body.decode("utf-8"))
        status = data.get("status")
        pregunta_id = request.GET.get("id")
        tipo = request.GET.get("tipo", "pre")

        # 2 = Documento listo para guardar, 6 = Guardado forzado
        if status in [2, 6]:
            download_url = data.get("url")
            pregunta = Pregunta.objects.get(id=pregunta_id)
            
            response = requests.get(download_url, verify=False, timeout=10)

            if response.status_code == 200:
                content_data = response.content
                
                if tipo == 'pre':
                    import re
                    doc = Document(io.BytesIO(content_data))
                    encontrado = False
                    
                    for p in doc.paragraphs:
                        if encontrado: break 
                        
                        for run in p.runs:
                            if run.font.highlight_color not in [None, WD_COLOR_INDEX.AUTO]:
                                texto_resaltado = run.text.strip().upper()
                                match = re.search(r'([A-E])(?:\s|[\)\.\-]|)', texto_resaltado)
                                
                                if match:
                                    pregunta.respuesta = match.group(1)
                                    encontrado = True
                                    break
                    
                    pregunta.contenido.save(f"{pregunta.nombre}.docx", ContentFile(content_data), save=False)
                else:
                    pregunta.solucion_archivo.save(f"sol_{pregunta.nombre}.docx", ContentFile(content_data), save=False)
                
                pregunta.save()

        return JsonResponse({"error": 0})
    except Exception as e:
        print(f"ERROR CALLBACK: {str(e)}")
        return JsonResponse({"error": 1, "message": str(e)})

@login_required
@role_required('admin', 'user')
def agregar_solucion_ajax(request, pregunta_id):
    pregunta = get_object_or_404(Pregunta, id=pregunta_id)
    if not pregunta.solucion_archivo:
        pregunta.solucion_archivo.save(
            f"sol_{pregunta.nombre}.docx", 
            ContentFile(crear_docx_minimo(ancho_cm=7)) 
        )
        pregunta.tiene_solucion = True
        pregunta.save()
    
    token, payload = generar_token_office(request, pregunta, es_solucion=True)
    return JsonResponse({'token': token, 'payload': payload})

@login_required
@role_required('admin', 'user')
def pregunta_edit(request, pregunta_id):
    """Reabre el editor de OnlyOffice con validación de autoría."""
    pregunta = get_object_or_404(Pregunta, id=pregunta_id)
    
    if not request.user.is_staff and pregunta.usuario != request.user.userprofile:
        raise PermissionDenied("No tienes permiso para editar esta pregunta.")

    token, payload = generar_token_office(request, pregunta, modo='edit')

    return render(request, 'Preguntas/preguntas/editor_continuo.html', {
        'pregunta': pregunta,
        'DOC_TOKEN': token,
        'PAYLOAD': payload,
        'ONLYOFFICE_API_URL': settings.ONLYOFFICE_API_URL,
        'es_edicion_manual': True
    })

@login_required
@role_required('admin', 'user')
def solucion_edit(request, pregunta_id):
    """Abre el editor de OnlyOffice para la solución."""
    pregunta = get_object_or_404(Pregunta, id=pregunta_id)
    
    if not request.user.is_staff and pregunta.usuario != request.user.userprofile:
        raise PermissionDenied("No tienes permiso.")

    # Si no existe, crear el docx de 14cm
    if not pregunta.solucion_archivo:
        pregunta.solucion_archivo.save(
            f"sol_{pregunta.nombre}.docx", 
            ContentFile(crear_docx_minimo(ancho_cm=7))
        )
        pregunta.tiene_solucion = True
        pregunta.save()

    token, payload = generar_token_office(request, pregunta, modo='edit', es_solucion=True)

    return render(request, 'Preguntas/preguntas/editor_continuo.html', {
        'pregunta': pregunta,
        'DOC_TOKEN': token,
        'PAYLOAD': payload,
        'ONLYOFFICE_API_URL': settings.ONLYOFFICE_API_URL,
        'es_edicion_manual': True,
        'es_solucion': True
    })

@login_required
@role_required('admin', 'user')
def generic_preview(request, pregunta_id, tipo='pregunta'):
    """Vista unificada para previsualizar Enunciados o Soluciones."""
    pregunta = get_object_or_404(Pregunta, id=pregunta_id)
    es_sol = (tipo == 'solucion')
    
    # Seleccionar el archivo correcto
    archivo = pregunta.solucion_archivo if es_sol else pregunta.contenido
    if not archivo:
        return JsonResponse({"error": f"No hay archivo de {tipo} disponible."}, status=404)

    # Generamos el token dinámico
    token, payload = generar_token_office(request, pregunta, modo='view', es_solucion=es_sol)

    return render(request, "Preguntas/preguntas/pregunta_preview.html", {
        "DOC_TOKEN": token,
        "PAYLOAD": payload,
        "pregunta": pregunta,
        "ONLYOFFICE_API_URL": settings.ONLYOFFICE_API_URL,
        "titulo_preview": "Vista Previa de Solución" if es_sol else "Vista Previa de Pregunta"
    })