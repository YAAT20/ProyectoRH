import random
from docx import Document
from ..models import *
from django.shortcuts import redirect
from django.shortcuts import render, get_object_or_404
from django.utils import timezone
from docx import Document
from docxcompose.composer import Composer
import os, jwt, time
from django.conf import settings
from django.contrib.auth.decorators import login_required
import math
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from django.db import transaction
from django.contrib import messages
import re
from docx.text.paragraph import Paragraph

def _limpiar_y_ajustar_documento(doc, quitar_resaltado=False):
    """
    Aplica interlineado 1.0, elimina todos los espacios y limpia 
    profundamente el resaltado (runs y párrafos).
    """
    def procesar_contenedor(contenedor):
        for paragraph in contenedor.paragraphs:
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            
            if quitar_resaltado:
                p_obj = paragraph._element
                pPr = p_obj.get_or_add_pPr()
                shd = pPr.xpath('./w:shd')
                if shd:
                    pPr.remove(shd[0])

                for run in paragraph.runs:
                    run.font.highlight_color = None
                    run._element.xpath('./w:rPr/w:shd')
                    rPr = run._element.get_or_add_rPr()
                    shd_run = rPr.xpath('./w:shd')
                    if shd_run:
                        rPr.remove(shd_run[0])

    procesar_contenedor(doc)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                procesar_contenedor(cell)

def _configurar_columnas_margenes(doc):
    """Aplica márgenes de 0.5cm y configura 3 columnas uniformes."""
    for section in doc.sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)

        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '3')
        cols.set(qn('w:space'), '708')
        cols.set(qn('w:equalWidth'), '1')

def generar_docx(practica, relaciones_preguntas):
    base = Document()
    # Eliminar párrafo inicial vacío
    if len(base.paragraphs) > 0:
        p = base.paragraphs[0]._element
        p.getparent().remove(p)
                  
    composer = Composer(base)

    for i, pp in enumerate(relaciones_preguntas, 1):
        p = pp.pregunta
        
        if not p.contenido or not os.path.exists(p.contenido.path):
            continue

        try:
            doc_temp = Document(p.contenido.path)
            
            # 1. Eliminar sección de solución dentro del Word de la pregunta
            body = doc_temp._element.body
            eliminar_resto = False
            for element in list(body):
                if eliminar_resto:
                    body.remove(element)
                    continue
                if element.tag.endswith('p'):
                    p_obj = Paragraph(element, doc_temp)
                    if re.search(r'@?SOLUCI[OÓ]N@?', p_obj.text, re.IGNORECASE):
                        eliminar_resto = True
                        body.remove(element)
            
            # 2. Limpiar formatos y resaltados
            _limpiar_y_ajustar_documento(doc_temp, quitar_resaltado=True)

            # 3. Insertar número y nombre de pregunta en NEGRITA
            if len(doc_temp.paragraphs) > 0:
                nombre_preg = p.nombre if p.nombre else f"ID-{p.id}"
                # Insertamos el párrafo al inicio
                p_num = doc_temp.paragraphs[0].insert_paragraph_before()
                run_num = p_num.add_run(f"{i}. {nombre_preg}")
                run_num.bold = True  # <--- NOMBRE EN NEGRITA
                
                p_num.paragraph_format.line_spacing = 1.0
                p_num.paragraph_format.space_before = Pt(0)
                p_num.paragraph_format.space_after = Pt(2) # Un mínimo respiro
            
            composer.append(doc_temp)

        except Exception as e:
            print(f"Error en pregunta {p.id}: {str(e)}")
            continue

    _configurar_columnas_margenes(base)

    carpeta = "media/practicas"
    os.makedirs(carpeta, exist_ok=True)
    path = f"{carpeta}/practica_{practica.uuid}.docx"
    composer.save(path)
    return path

def generar_docx_solucionario(practica):
    preguntas_practica = practica.preguntas.select_related('pregunta').all().order_by('id')
    
    base = Document()
    if len(base.paragraphs) > 0:
        p = base.paragraphs[0]._element
        p.getparent().remove(p)
          
    composer = Composer(base)

    for i, pp in enumerate(preguntas_practica, 1):
        try:
            if pp.pregunta.contenido and os.path.exists(pp.pregunta.contenido.path):
                doc_enunciado = Document(pp.pregunta.contenido.path)
                _limpiar_y_ajustar_documento(doc_enunciado, quitar_resaltado=False)

                if len(doc_enunciado.paragraphs) > 0:
                    nombre_preg = pp.pregunta.nombre if pp.pregunta.nombre else f"ID-{pp.pregunta.id}"
                    p_num = doc_enunciado.paragraphs[0].insert_paragraph_before()
                    run_num = p_num.add_run(f"{i}. {nombre_preg}")
                    run_num.bold = True
                    p_num.paragraph_format.space_before = Pt(10)
                
                composer.append(doc_enunciado)

            if pp.pregunta.solucion_archivo and os.path.exists(pp.pregunta.solucion_archivo.path):
                doc_sol = Document(pp.pregunta.solucion_archivo.path)
                _limpiar_y_ajustar_documento(doc_sol, quitar_resaltado=False)
                
                if len(doc_sol.paragraphs) > 0:
                    p_label = doc_sol.paragraphs[0].insert_paragraph_before()
                    run_label = p_label.add_run("SOLUCIÓN:")
                    run_label.bold = True
                
                composer.append(doc_sol)
            
        except Exception as e:
            continue

    _configurar_columnas_margenes(base)
    path = f"media/practicas/solucionario_{practica.uuid}.docx"
    composer.save(path)
    return path

def tiempo_agotado(practica):
    segundos_totales = practica.cantidad_preguntas * 90
    limite = practica.fecha_inicio + timedelta(seconds=segundos_totales)
    return timezone.now() > limite

@login_required
def crear_practica(request):
    if request.method == "POST":
        curso_id = request.POST.get("curso")
        temas_ids = request.POST.getlist("temas")
        cantidad = int(request.POST.get("cantidad") or 0)
        usuario_actual = request.user.userprofile
        
        practicas_anteriores = Practica.objects.filter(usuario=usuario_actual)
        for p_vieja in practicas_anteriores:
            rutas_a_borrar = [
                f"media/practicas/solucionario_{p_vieja.uuid}.docx",
                f"media/practicas/practica_{p_vieja.uuid}.docx" 
            ]
            for ruta in rutas_a_borrar:
                if os.path.exists(ruta):
                    try: os.remove(ruta)
                    except: pass
                        
        preguntas_ids = list(Pregunta.objects.filter(
            curso_id=curso_id,
            tema_id__in=temas_ids
        ).values_list('id', flat=True))

        if len(preguntas_ids) < cantidad:
            messages.error(request, "No hay suficientes preguntas.")
            return redirect("crear_practica")

        ids_seleccionados = random.sample(preguntas_ids, cantidad)
        tiempo_calculado = math.ceil(cantidad * 1.5)
        practica = None

        try:
            with transaction.atomic():
                practica = Practica.objects.create(
                    usuario=usuario_actual,
                    curso_id=curso_id,
                    cantidad_preguntas=cantidad,
                    tiempo_minutos=tiempo_calculado,
                    finalizado=False
                )
                practica.temas.set(temas_ids)

                PracticaPregunta.objects.bulk_create([
                    PracticaPregunta(practica=practica, pregunta_id=pid) 
                    for pid in ids_seleccionados
                ])

            relaciones = practica.preguntas.all().order_by('id').select_related('pregunta')
            
            generar_docx(practica, relaciones)

            return redirect("resolver_practica", practica.uuid)

        except Exception as e:
            if practica: practica.delete()
            messages.error(request, f"Error: {str(e)}")
            return redirect("crear_practica")

    universidades = Universidad.objects.all().order_by('nombre')
    return render(request, "Preguntas/practicas/crear_practica.html", {"universidades": universidades})

def generar_token_practica(practica, modo='view', es_solucionario=False):
    prefijo = "solucionario" if es_solucionario else "practica"
    ip_interna = "http://192.168.18.20:8003"
    
    file_url = f"{ip_interna}/banco/media/practicas/{prefijo}_{practica.uuid}.docx"
    path = f"media/practicas/{prefijo}_{practica.uuid}.docx"

    try:
        version_key = int(os.path.getmtime(path))
    except Exception:
        version_key = practica.uuid

    document_key = f"{prefijo.upper()}_{practica.uuid}_{version_key}"

    payload = {
        "iat": int(time.time()),
        "exp": int(time.time()) + 3600,
        "document": {
            "fileType": "docx",
            "key": document_key,
            "url": file_url,
            "permissions": {
                "edit": False,
                "download": False,
                "copy": False,
                "ai": False,
                "print": False,
                "fillForms": False,
                "comment": False,
                "modifyFilter": False,
                "modifyContentControl": False
            }
        },
        "editorConfig": {
            "mode": modo,
            "lang": "es",
            "customization": {
                "chat": False,
                "help": False,
                "comments": False,
                "search": False,
                "compactHeader": True,
                "toolbarNoTabs": True,
                "hideRightMenu": True,
                "autosave": False,
                "toolbar": not es_solucionario,
                "forcesave": False,
                "customer": {
                    "name": "Academia Robert Hooke",
                    "address": "Cajamarca, Perú",
                },
                "plugins": False
            }
        }
    }

    token = jwt.encode(payload, settings.ONLYOFFICE_JWT_SECRET, algorithm="HS256")
    payload["token"] = token

    return payload

@login_required
def resolver_practica(request, practica_uuid):
    with transaction.atomic():
        practica = get_object_or_404(
            Practica.objects.select_for_update(), 
            uuid=practica_uuid, 
            usuario=request.user.userprofile
        )

        ruta_docx = f"media/practicas/practica_{practica.uuid}.docx"

        def limpiar_docx_temporal():
            if os.path.exists(ruta_docx):
                try:
                    os.remove(ruta_docx)
                except Exception as e:
                    print(f"Error al eliminar archivo: {e}")

        debe_finalizar = False

        if practica.finalizado or tiempo_agotado(practica):
            debe_finalizar = True

        elif request.method == "POST":
            v_aciertos = v_errores = v_blancos = 0
            
            preguntas_practica = practica.preguntas.select_related('pregunta').all().order_by('id')
            preguntas_a_actualizar = []
            
            for pp in preguntas_practica:
                resp = request.POST.get(f"pregunta_{pp.id}")
                pp.respuesta_alumno = resp
                preguntas_a_actualizar.append(pp)

                if not resp:
                    v_blancos += 1
                elif resp == pp.pregunta.respuesta:
                    v_aciertos += 1
                else:
                    v_errores += 1

            if preguntas_a_actualizar:
                PracticaPregunta.objects.bulk_update(preguntas_a_actualizar, ['respuesta_alumno'])

            practica.aciertos = v_aciertos
            practica.errores = v_errores
            practica.blancos = v_blancos
            debe_finalizar = True

        if debe_finalizar:
            if not practica.finalizado:
                practica.finalizado = True
                practica.fecha_fin = timezone.now()
                practica.save()
            
            transaction.on_commit(lambda p=practica: generar_docx_solucionario(p))
            transaction.on_commit(limpiar_docx_temporal)
            
            return redirect("resultado_practica", practica.uuid)

    config = generar_token_practica(practica)
    preguntas_ordenadas = practica.preguntas.all().order_by('id')
    
    return render(request, "Preguntas/practicas/examen.html", {
        "practica": practica,
        "preguntas": preguntas_ordenadas,
        "config": config
    })

@login_required
def resultado_practica(request, practica_uuid):
    practica = get_object_or_404(
        Practica, 
        uuid=practica_uuid,
        usuario=request.user.userprofile
    )
    if not practica.finalizado:
        messages.warning(request, "Aún no has finalizado este simulacro.")
        return redirect("resolver_practica", practica.uuid)
    
    total = practica.cantidad_preguntas
    
    def get_pct(val):
        return round((val / total) * 100) if total > 0 else 0

    pct_aciertos = get_pct(practica.aciertos)
    pct_errores = get_pct(practica.errores)
    pct_blancos = get_pct(practica.blancos)

    suma_pct = pct_aciertos + pct_errores + pct_blancos
    if suma_pct != 100 and total > 0:
        pct_aciertos += (100 - suma_pct)

    detalles = []
    
    preguntas_qs = PracticaPregunta.objects.filter(
        practica=practica
    ).select_related('pregunta').order_by('id')
    
    for i, pp in enumerate(preguntas_qs, 1):
        if not pp.respuesta_alumno:
            estado = "blanco"
        elif pp.respuesta_alumno == pp.pregunta.respuesta:
            estado = "correcto"
        else:
            estado = "incorrecto"

        detalles.append({
            "numero": i,
            "nombre": pp.pregunta.nombre,
            "correcta": pp.pregunta.respuesta,
            "marcada": pp.respuesta_alumno,
            "estado": estado
        })

    config_solucionario = generar_token_practica(practica, es_solucionario=True)

    return render(request, "Preguntas/practicas/resultado_practica.html", {
        "practica": practica,
        "aciertos": practica.aciertos,
        "errores": practica.errores,
        "blancos": practica.blancos,
        "total": total,
        "pct_aciertos": pct_aciertos,
        "pct_errores": pct_errores,
        "pct_blancos": pct_blancos,
        "detalles": detalles,
        "config_solucionario": config_solucionario
    })