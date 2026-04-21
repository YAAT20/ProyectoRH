import os
import tempfile
from contextlib import contextmanager
from io import BytesIO
from copy import deepcopy
from docx import Document
from django.contrib.auth.decorators import login_required
from django.contrib.admin.views.decorators import staff_member_required
from django.contrib import messages
from django.shortcuts import render, redirect
from django.core.files import File
from ..forms import *
from .auth_views import exclude_supervisor, role_required
from ..models import Pregunta, UserProfile
from docx.oxml.ns import qn
from copy import deepcopy
import io
import logging

@contextmanager
def temp_docx_file(content_bytes, suffix='.docx'):
    """Context manager mejorado para archivos temporales"""
    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
            temp_file.write(content_bytes)
            temp_path = temp_file.name
        yield temp_path
    finally:
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except:
                pass

def copy_related_parts(new_doc, original_doc, block_elements):
    """
    Copia las partes relacionadas (imágenes, gráficos) que son referenciadas
    en los elementos del bloque al nuevo documento.
    """
    original_rels = original_doc.part.rels
    
    used_rIds = set()
    for element in block_elements:
        for rel_attr in ['r:embed', 'r:link', 'v:imagedata']:
            for el in element.xpath(f'.//@{rel_attr}'):
                used_rIds.add(el)
    
    for rId in used_rIds:
        if rId in original_rels:
            rel = original_rels[rId]
            if rel.is_external:
                new_doc.part.rels.add_relationship(
                    rel.reltype,
                    rel._target,
                    rel.rId,
                    is_external=True
                )
            else:
                new_doc.part.rels.add_relationship(
                    rel.reltype,
                    rel.target_part,
                    rel.rId
                )

def create_exact_copy_docx(original_doc, block_elements):
    """
    Crea un nuevo documento que conserva exactamente todo el contenido
    incluyendo imágenes, tablas, ecuaciones, con sus formatos originales.
    Versión corregida del error 'element has no setter'.
    """
    new_doc = Document()
    
    for style in original_doc.styles:
        try:
            if style.name not in new_doc.styles:
                new_style = new_doc.styles.add_style(style.name, style.type)
                if style.font:
                    new_style.font.name = style.font.name
                    new_style.font.size = style.font.size
                    new_style.font.bold = style.font.bold
                    new_style.font.italic = style.font.italic
                    new_style.font.underline = style.font.underline
                if hasattr(style, 'paragraph_format'):
                    new_style.paragraph_format.alignment = style.paragraph_format.alignment
                    new_style.paragraph_format.left_indent = style.paragraph_format.left_indent
                    new_style.paragraph_format.right_indent = style.paragraph_format.right_indent
        except Exception as e:
            print(f"Warning: No se pudo copiar el estilo {style.name}: {str(e)}")
            continue
    
    new_doc._element.body.clear_content()
    
    if original_doc.sections:
        new_sect = new_doc.sections[0]
        orig_sect = original_doc.sections[0]
        new_sect.start_type = orig_sect.start_type
        new_sect.orientation = orig_sect.orientation
        new_sect.page_width = orig_sect.page_width
        new_sect.page_height = orig_sect.page_height
        new_sect.left_margin = orig_sect.left_margin
        new_sect.right_margin = orig_sect.right_margin
        new_sect.top_margin = orig_sect.top_margin
        new_sect.bottom_margin = orig_sect.bottom_margin
        new_sect.header_distance = orig_sect.header_distance
        new_sect.footer_distance = orig_sect.footer_distance
    
    new_body = new_doc._element.body
    for element in block_elements:
        new_body.append(deepcopy(element))
    
    copy_related_parts(new_doc, original_doc, block_elements)
    
    return new_doc

def detectar_clave_resaltada(elements):
    """
    Analiza los elementos del bloque buscando texto con resaltado amarillo.
    Retorna 'A', 'B', 'C', 'D' o 'E' si lo encuentra.
    """
    for element in elements:
        runs = element.xpath('.//w:r')
        for r in runs:
            rPr = r.find(qn('w:rPr'))
            if rPr is not None:
                highlight = rPr.find(qn('w:highlight'))
                if highlight is not None and highlight.get(qn('w:val')) == 'yellow':
                    texto = "".join(t.text for t in r.xpath('.//w:t') if t.text).strip().upper()
                    for letra in ['A', 'B', 'C', 'D', 'E']:
                        if letra in texto:
                            return letra
    return None

logger = logging.getLogger(__name__)

@login_required
@staff_member_required
@exclude_supervisor
@role_required('admin')
def masivo_pregunta_create(request):
    if request.method == 'POST':
        form = CargaMasivaPreguntaForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                archivo_word = request.FILES['archivo']
                archivo_bytes = archivo_word.read()
                user_profile = UserProfile.objects.get(user=request.user)
                
                with temp_docx_file(archivo_bytes) as temp_path:
                    doc = Document(temp_path)
                    all_blocks, current_block = [], []
                    
                    for element in doc.element.body:
                        txt = ''.join(t.text for t in element.xpath('.//w:t') if t.text).strip()
                        if txt == '*****':
                            if current_block: 
                                all_blocks.append(current_block)
                                current_block = []
                        else:
                            current_block.append(element)
                    if current_block: all_blocks.append(current_block)

                for i, block in enumerate(all_blocks, start=1):
                    clave = detectar_clave_resaltada(block) or form.cleaned_data['respuesta_default']
                    enunciado_els, solucion_els = [], []
                    encontrado_sol = False
                    
                    for el in block:
                        el_txt = ''.join(t.text for t in el.xpath('.//w:t') if t.text).strip().upper()
                        if '@SOLUCIÓN@' in el_txt:
                            encontrado_sol = True
                            continue
                        (solucion_els if encontrado_sol else enunciado_els).append(el)

                    with temp_docx_file(archivo_bytes) as temp_path:
                        orig_doc = Document(temp_path)
                        
                        doc_p = create_exact_copy_docx(orig_doc, enunciado_els)
                        buf_p = io.BytesIO()
                        doc_p.save(buf_p)
                        buf_p.seek(0)

                        buf_s = None
                        if encontrado_sol and solucion_els:
                            doc_s = create_exact_copy_docx(orig_doc, solucion_els)
                            buf_s = io.BytesIO()
                            doc_s.save(buf_s)
                            buf_s.seek(0)

                    preg = Pregunta(
                        universidad=form.cleaned_data['universidad'],
                        curso=form.cleaned_data['curso'],
                        tema=form.cleaned_data['tema'],
                        nivel=form.cleaned_data['nivel'],
                        respuesta=clave,
                        usuario=user_profile,
                        tiene_solucion=encontrado_sol
                    )
                    preg.save() 
                    
                    if buf_p:
                        preg.contenido.save(f"P_{preg.nombre}.docx", File(buf_p), save=False)
                    
                    if buf_s:
                        preg.solucion_archivo.save(f"S_{preg.nombre}.docx", File(buf_s), save=False)
                    
                    preg.save()
                    
                    buf_p.close()
                    if buf_s: buf_s.close()

                messages.success(request, f'Éxito: Se cargaron {len(all_blocks)} preguntas correctamente.')
                return redirect('pregunta-list')

            except Exception as e:
                logger.error(f"Error masivo crítico: {e}", exc_info=True)
                messages.error(request, f"Error durante el procesamiento: {e}")
        
        return render(request, 'Preguntas/preguntas/masivo_pregunta_form.html', {'form': form})

    return render(request, 'Preguntas/preguntas/masivo_pregunta_form.html', {'form': CargaMasivaPreguntaForm()})