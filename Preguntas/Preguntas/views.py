# Importaciones organizadas por grupos y eliminadas las no utilizadas
import base64
import io, os
from datetime import timedelta
import mammoth
from django.http import JsonResponse
import subprocess
import csv
from docx.oxml import OxmlElement, ns
from docx.shared import Pt, Inches  # Para tamaño de fuente y los márgenes del doc final xd 
from collections import defaultdict
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pdfkit
from django.conf import settings
from django.core.files.storage import default_storage

# Django
from django.shortcuts import render, redirect, get_object_or_404
from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy
from django.contrib import messages
from django.http import JsonResponse, HttpResponse
from django.utils import timezone
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import UserPassesTestMixin, LoginRequiredMixin
from django.core.exceptions import PermissionDenied
from django.contrib.auth.models import User
from django.contrib.auth.forms import UserCreationForm
from django.contrib.admin.views.decorators import staff_member_required
from django.core.paginator import Paginator

# Librerías de terceros
from docx import Document
from docxcompose.composer import Composer
try:
    from docxcompose.composer import ImportFormatMode
except ImportError:
    ImportFormatMode = None

# Aplicación
from .models import Universidad, Curso, Tema, Pregunta
from .forms import UniversidadForm, CursoForm, TemaForm, PreguntaForm, FiltroPreguntaForm
from django.contrib import admin
from .models import UserProfile


# Página principal
@login_required
def home(request):
    return render(request, 'Preguntas/home.html')

class AdminRequiredMixin(LoginRequiredMixin, UserPassesTestMixin):    
    def test_func(self):
        return self.request.user.is_staff  # Solo administradores pueden acceder
    
    def handle_no_permission(self):
        messages.error(self.request, "No tienes permiso para realizar esta acción.")
        raise PermissionDenied

# Mixin para mensajes de éxito en vistas basadas en clases
class SuccessMessageMixin:
    success_message = ""
    
    def form_valid(self, form):
        messages.success(self.request, self.success_message)
        return super().form_valid(form)
    
    def delete(self, request, *args, **kwargs):
        messages.success(self.request, self.success_message)
        return super().delete(request, *args, **kwargs)


# CRUD Universidades
class UniversidadListView(LoginRequiredMixin, ListView):
    model = Universidad
    template_name = 'Preguntas/universidad_list.html'
    context_object_name = 'universidades'


class UniversidadCreateView(LoginRequiredMixin, SuccessMessageMixin, CreateView):
    model = Universidad
    form_class = UniversidadForm
    template_name = 'Preguntas/universidad_form.html'
    success_url = reverse_lazy('universidad-list')
    success_message = 'Universidad creada exitosamente.'


class UniversidadUpdateView(AdminRequiredMixin, SuccessMessageMixin, UpdateView):
    model = Universidad
    form_class = UniversidadForm
    template_name = 'Preguntas/universidad_form.html'
    success_url = reverse_lazy('universidad-list')
    success_message = 'Universidad actualizada exitosamente.'


class UniversidadDeleteView(AdminRequiredMixin, SuccessMessageMixin, DeleteView):
    model = Universidad
    template_name = 'Preguntas/universidad_confirm_delete.html'
    success_url = reverse_lazy('universidad-list')
    success_message = 'Universidad eliminada exitosamente.'


# CRUD Cursos
class CursoListView(LoginRequiredMixin, ListView):
    model = Curso
    template_name = 'Preguntas/curso_list.html'
    context_object_name = 'cursos'

    def get_queryset(self):
        queryset = super().get_queryset()
        universidad_id = self.request.GET.get('universidad')
        if universidad_id:
            queryset = queryset.filter(universidad_id=universidad_id)
        return queryset
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context.update({
            'universidades': Universidad.objects.all(),
            'universidad_id': self.request.GET.get('universidad')
        })
        return context


class CursoCreateView(LoginRequiredMixin, SuccessMessageMixin, CreateView):
    model = Curso
    form_class = CursoForm
    template_name = 'Preguntas/curso_form.html'
    success_url = reverse_lazy('curso-list')
    success_message = 'Curso creado exitosamente.'

    def get_initial(self):
        initial = super().get_initial()
        universidad_id = self.request.GET.get('universidad_id')
        if universidad_id:
            initial['universidad'] = Universidad.objects.get(id=universidad_id)
        return initial


class CursoUpdateView(LoginRequiredMixin, SuccessMessageMixin, UpdateView):
    model = Curso
    form_class = CursoForm
    template_name = 'Preguntas/curso_form.html'
    success_url = reverse_lazy('curso-list')
    success_message = 'Curso actualizado exitosamente.'


class CursoDeleteView(AdminRequiredMixin, SuccessMessageMixin, DeleteView):
    model = Curso
    template_name = 'Preguntas/curso_confirm_delete.html'
    success_url = reverse_lazy('curso-list')
    success_message = 'Curso eliminado exitosamente.'   


# CRUD Temas
class TemaListView(LoginRequiredMixin, ListView):
    model = Tema
    template_name = 'Preguntas/tema_list.html'
    context_object_name = 'temas'

    def get_queryset(self):
        queryset = super().get_queryset().select_related('curso__universidad')
        filtros = {}

        curso_id = self.request.GET.get('curso')
        universidad_id = self.request.GET.get('universidad')

        if curso_id:
            filtros['curso_id'] = curso_id
        if universidad_id:
            filtros['curso__universidad_id'] = universidad_id

        return queryset.filter(**filtros) if filtros else queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context.update({
            'universidades': Universidad.objects.all(),
            'cursos': Curso.objects.all(),
            'universidad_id': self.request.GET.get('universidad'),
            'curso_id': self.request.GET.get('curso')
        })
        return context


class TemaCreateView(LoginRequiredMixin, SuccessMessageMixin, CreateView):
    model = Tema
    form_class = TemaForm
    template_name = 'Preguntas/tema_form.html'
    success_url = reverse_lazy('tema-list')
    success_message = 'Tema creado exitosamente.'

    def get_initial(self):
        initial = super().get_initial()
        curso_id = self.request.GET.get('curso_id')
        if curso_id:
            initial['curso'] = Curso.objects.get(id=curso_id)
        return initial

class TemaUpdateView(LoginRequiredMixin, SuccessMessageMixin, UpdateView):
    model = Tema
    form_class = TemaForm
    template_name = 'Preguntas/tema_form.html'
    success_url = reverse_lazy('tema-list')
    success_message = 'Tema actualizado exitosamente.'


class TemaDeleteView(AdminRequiredMixin, SuccessMessageMixin, DeleteView):
    model = Tema
    template_name = 'Preguntas/tema_confirm_delete.html'
    success_url = reverse_lazy('tema-list')
    success_message = 'Tema eliminado exitosamente.'


def extract_equations(doc_path):
    doc = Document(doc_path)
    equations = []
    for para in doc.paragraphs:
        # Busca el primer elemento "oMath" en el párrafo
        if para._element.xpath('.//w:oMath'):
            equations.append(para.text)  # O extrae el contenido de la ecuación
    return equations

# Función para convertir DOCX a HTML
def document_to_html(doc_path):
    try:
        doc = Document(doc_path)
        html_content = []

        # Procesar párrafos y títulos
        for para in doc.paragraphs:
            if para.style and para.style.name.startswith("Heading"):
                level = ''.join(filter(str.isdigit, para.style.name))
                level = int(level) if level else 2
                html_content.append(f"<h{level}>{para.text}</h{level}>")
            else:
                html_content.append(f"<p>{para.text}</p>")

        # Procesar tablas y ecuaciones
        for table in doc.tables:
            table_html = ["<table style='width: 100%; border: 1px solid black; margin: 10px 0;'>"]
            for row in table.rows:
                table_html.append("<tr>")
                for cell in row.cells:
                    cell_text = "<br>".join(p.text for p in cell.paragraphs)
                    cell_eq = extract_equations(cell)
                    table_html.append(f"<td style='padding: 5px;'>{cell_text}{cell_eq}</td>")
                table_html.append("</tr>")
            table_html.append("</table>")
            html_content.append("".join(table_html))

        # Procesar imágenes
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                img_part = rel.target_part
                img_bytes = img_part.blob
                img_base64 = base64.b64encode(img_bytes).decode('utf-8')
                img_mime = "image/png"
                html_content.append(
                    f'<img src="data:{img_mime};base64,{img_base64}" '
                    f'alt="Imagen del documento" style="max-width: 100%; margin: 10px 0;">'
                )

        return "".join(html_content)

    except Exception as e:
        return f"<p>Error al procesar el documento: {str(e)}</p>"


#desde aquí empecé a modificar lo del formato de las preguntas
#para darle 2 columnas al doc final
def set_tres_columns(section):
    sectPr = section._sectPr  # Obtener el elemento de la sección
    cols = OxmlElement('w:cols')
    cols.set(ns.qn('w:num'), '3')  # Establecer dos columnas
    sectPr.append(cols)

def set_margenes(section):
    """Configura los márgenes del documento según lo solicitado."""
    section.top_margin = Inches(2 / 2.54)  # 2 cm
    section.left_margin = Inches(0.76 / 2.54)  # 0.76 cm
    section.right_margin = Inches(0.76 / 2.54)  # 0.76 cm
    section.bottom_margin = Inches(3.25 / 2.54)  # 3.25 cm

def aplicar_formato_texto(doc):
    """Aplica Arial Narrow y tamaño 9 pt a todo el contenido del documento."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial Narrow"
            run.font.size = Pt(9)
            r = run._element
            r.rPr.rFonts.set(qn("w:eastAsia"), "Arial Narrow")
    for style in doc.styles:
        if style.type == 1:  # Solo afecta estilos de párrafo
            if style.name.lower() in ["list paragraph", "lista numerada", "lista con viñetas"]:
                style.font.name = "Arial Narrow"
                style.font.size = Pt(9)

def combinar_documentos(preguntas):
    """Combina documentos de preguntas agrupadas por curso y tema en formato de tres columnas."""
    master_doc = Document()
    master_doc.add_heading("Preguntas Combinadas", level=1)
    composer = Composer(master_doc)
    
    # Configurar el documento con tres columnas y los margenes caumsa
    set_tres_columns(master_doc.sections[0])

    set_margenes(master_doc.sections[0])

    # Agrupar preguntas por curso y tema
    preguntas_ordenadas = defaultdict(lambda: defaultdict(list))

    for pregunta in preguntas:
        if pregunta.contenido and hasattr(pregunta.contenido, 'path'):
            preguntas_ordenadas[pregunta.curso.nombre][pregunta.tema.nombre].append(pregunta)

    # Iterar sobre cursos y temas
    for curso, temas in sorted(preguntas_ordenadas.items()):
        master_doc.add_heading(f"Curso: {curso}", level=1)

        for tema, preguntas_tema in sorted(temas.items()):
            master_doc.add_heading(f"Tema: {tema}", level=2)

            for pregunta in preguntas_tema:
                master_doc.add_heading(f"Pregunta: {pregunta.nombre}", level=3)
                sub_doc = Document(pregunta.contenido.path)

                # Aplicar formato de texto en Arial Narrow 9 pt al subdocumento
                aplicar_formato_texto(sub_doc)

                if ImportFormatMode is not None:
                    composer.append(sub_doc, import_format=ImportFormatMode.KEEP_SOURCE_FORMATTING)
                else:
                    composer.append(sub_doc)

    # Aplicar formato de texto en el documento final
    aplicar_formato_texto(master_doc)

    # Guardar el documento en memoria
    buffer = io.BytesIO()
    composer.save(buffer)
    buffer.seek(0)
    
    return buffer


# Gestión de Preguntas
@login_required
def pregunta_list(request):
    # Obtener el perfil del usuario actual
    user_profile = get_object_or_404(UserProfile, user=request.user)

    total_preguntas = Pregunta.objects.filter(usuario=user_profile).count()


    # Filtrar preguntas según permisos
    if request.user.is_superuser:
        preguntas = Pregunta.objects.filter(usuario=user_profile)  # Cambiar a user_profile
    else:
        fecha_limite = timezone.now() - timedelta(days=1)
        preguntas = Pregunta.objects.filter(
            usuario=user_profile,  # Cambiar a user_profile
            fecha_creacion__gte=fecha_limite
        )

    # Aplicar filtros
    filtros = {}
    form = FiltroPreguntaForm(request.GET or None)
    
    for campo in ['universidad', 'curso', 'tema', 'nivel']:
        valor = request.GET.get(campo)
        if valor:
            filtros[campo + '_id' if campo != 'nivel' else campo] = valor
    
    preguntas = preguntas.filter(**filtros) if filtros else preguntas

    context = {
        'total_preguntas': total_preguntas,  # Contador de preguntas 
        'preguntas': preguntas,
        'form': form,
        'universidades': Universidad.objects.all(),
        'cursos': Curso.objects.all(),
        'temas': Tema.objects.all(),
    }
    
    # Agregar los valores de filtro al contexto
    for campo in ['universidad', 'curso', 'tema', 'nivel']:
        context[f'{campo}_id'] = request.GET.get(campo)

    return render(request, 'Preguntas/pregunta_list.html', context)

@login_required
def pregunta_create(request):
    if request.method == 'POST':
        form = PreguntaForm(request.POST, request.FILES)
        if form.is_valid():
            pregunta = form.save(commit=False)
            user_profile = UserProfile.objects.get(user=request.user)
            pregunta.usuario = user_profile

            # Asignar nombre solo si se solicita
            if form.cleaned_data['add_nombre']:
                pregunta.nombre = form.cleaned_data['nombre']
            else:
                pregunta.nombre = None
            
            pregunta.save()
            
            # Guardar el archivo si se proporcionó
            if 'contenido' in request.FILES:
                pregunta.contenido = request.FILES['contenido']
                pregunta.save()
            
            messages.success(request, 'Pregunta creada exitosamente.')
            return redirect('pregunta-create')
    else:
        form = PreguntaForm()

    return render(request, 'Preguntas/pregunta_form.html', {
        'form': form,
        'title': 'Nueva Pregunta'
    })


@login_required
def pregunta_update(request, pk):
    # Obtener el perfil del usuario actual
    user_profile = get_object_or_404(UserProfile, user=request.user)
    
    # Cambiar la consulta para usar el perfil del usuario
    pregunta = get_object_or_404(Pregunta, pk=pk, usuario=user_profile)
    
    if request.method == 'POST':
        form = PreguntaForm(request.POST, request.FILES, instance=pregunta)
        if form.is_valid():
            form.save()
            messages.success(request, 'Pregunta actualizada exitosamente.')
            return redirect('pregunta-list')
    else:
        form = PreguntaForm(instance=pregunta)

    return render(request, 'Preguntas/pregunta_form.html', {
        'form': form,
        'pregunta': pregunta,
        'title': 'Editar Pregunta'
    })

@login_required
def pregunta_delete(request, pk):
    # Obtener el perfil del usuario actual
    user_profile = get_object_or_404(UserProfile, user=request.user)
    
    # Cambiar la consulta para usar el perfil del usuario
    pregunta = get_object_or_404(Pregunta, pk=pk, usuario=user_profile)
    
    if request.method == 'POST':
        pregunta.delete()
        messages.success(request, 'Pregunta eliminada exitosamente.')
        return redirect('pregunta-list')

    return render(request, 'Preguntas/pregunta_confirm_delete.html', {
        'pregunta': pregunta
    })

@login_required
def extract_text_from_docx(doc_path):
    """
    Extrae solo el texto del documento DOCX.
    Devuelve el contenido como una cadena HTML con párrafos.
    """
    doc = Document(doc_path)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(f"<p>{para.text}</p>")
    return "".join(paragraphs) if paragraphs else "<p>No hay contenido disponible.</p>"

def convert_image_callback(image):
    try:
        with image.open() as image_file:
            image_bytes = image_file.read()
            encoded = base64.b64encode(image_bytes).decode('ascii')
            return {"src": f"data:{image.content_type};base64,{encoded}"}
    except Exception as e:
        # Registra el error y devuelve una cadena vacía para la imagen
        print("Error processing image:", e)
        return {"src": ""}

@login_required
def vista_previa(request, pk):
    pregunta = get_object_or_404(Pregunta, pk=pk)

    # Verificar permisos
    if not request.user.is_staff and pregunta.usuario != request.user:
        return JsonResponse({'error': 'Acceso no autorizado'}, status=403)

    docx_text = "<p>No hay contenido disponible.</p>"
    if pregunta.contenido:
        try:
            file_path = pregunta.contenido.path if hasattr(pregunta.contenido, 'path') else pregunta.contenido
            ext = os.path.splitext(file_path)[1].lower()

            if ext == ".doc":
                converted_path = os.path.splitext(file_path)[0] + ".docx"
                if not os.path.exists(converted_path) or os.path.getmtime(file_path) > os.path.getmtime(converted_path):
                    subprocess.run(
                        [
                            "libreoffice",
                            "--headless",
                            "--convert-to",
                            "docx",
                            file_path,
                            "--outdir",
                            os.path.dirname(file_path)
                        ],
                        check=True
                    )
                file_path = converted_path

            if os.path.splitext(file_path)[1].lower() != ".docx":
                raise Exception("Formato no soportado para vista previa (se requiere DOCX).")

            # Convertir el documento a HTML
            with open(file_path, "rb") as docx_file:
                result = mammoth.convert_to_html(
                    docx_file,
                    convert_image=mammoth.images.img_element(convert_image_callback)
                )
                docx_text = result.value
                if result.messages:
                    print(f"Advertencias al convertir {file_path}: {result.messages}")

            # Extraer ecuaciones (si es necesario)
            equations = extract_equations(file_path)  # Asegúrate de que esta función esté bien implementada
            for eq in equations:
                docx_text += f"<p>$$ {eq} $$</p>"  # Asegúrate de que MathJax pueda procesar esto

        except subprocess.CalledProcessError as cpe:
            docx_text = f"<p>Error al convertir DOC a DOCX: {str(cpe)}</p>"
        except Exception as e:
            docx_text = f"<p>Error al procesar el documento: {str(e)}</p>"

    return JsonResponse({
        'pregunta_nombre': pregunta.nombre,
        'docx_text': docx_text
    })

@login_required
def descargar_preguntas(request):
    pregunta_ids = request.POST.getlist('preguntas')
    preguntas = Pregunta.objects.filter(id__in=pregunta_ids, usuario=request.user)
    
    if not preguntas:
        messages.error(request, 'No se encontraron preguntas para descargar.')
        return redirect('pregunta-list')
    
    buffer = combinar_documentos(preguntas)
    
    response = HttpResponse(
        buffer,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename="preguntas_combinadas.docx"'
    
    return response

# AJAX
@login_required
def load_cursos(request):
    universidad_id = request.GET.get('universidad_id')
    cursos = Curso.objects.filter(universidad_id=universidad_id)
    return JsonResponse([{'id': curso.id, 'nombre': str(curso)} for curso in cursos], safe=False)


@login_required
def load_temas(request):
    curso_id = request.GET.get('curso_id')
    temas = Tema.objects.filter(curso_id=curso_id)
    return JsonResponse([{'id': tema.id, 'nombre': str(tema)} for tema in temas], safe=False)


# Autenticación
def register(request):
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            messages.success(request, 'Registro exitoso. ¡Bienvenido!')
            return redirect('pregunta-list')
    else:
        form = UserCreationForm()
    return render(request, 'registration/register.html', {'form': form})


def user_login(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)

        if user is not None:
            try:
                user_profile = user.profile
            except UserProfile.DoesNotExist:
                user_profile = UserProfile.objects.create(user=user)

            if user_profile.is_active:
                login(request, user)
                next_url = request.GET.get('next', 'pregunta-list')
                return redirect(next_url)
            else:
                messages.error(request, 'Tu cuenta está suspendida. Comunícate con el administrador.')
                logout(request)
                return redirect('login')
        else:
            messages.error(request, 'Usuario o contraseña incorrectos.')
    
    return render(request, 'registration/login.html')

class UserProfileAdmin(admin.ModelAdmin):
    list_display = ('user', 'is_active')
    list_filter = ('is_active',)

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        return qs.select_related('user')

admin.site.register(UserProfile, UserProfileAdmin)

def user_logout(request):
    logout(request)
    messages.info(request, 'Has cerrado sesión correctamente.')
    return redirect('login')

@staff_member_required
def admin_dashboard(request):
    # Filtros
    filtros = {}
    for campo in ['tema', 'universidad', 'curso']:
        valor = request.GET.get(campo)
        if valor:
            filtros[campo + '__id'] = valor
    
    # Preguntas filtradas y ordenadas
    preguntas_qs = Pregunta.objects.filter(**filtros).order_by('-fecha_creacion')
    
    # Paginación
    paginator = Paginator(preguntas_qs, 20)
    page_number = request.GET.get('page')
    preguntas_recientes = paginator.get_page(page_number)
    
    # Estadísticas
    preguntas_por_usuario = {}
    for user in User.objects.all():
        preguntas_count = Pregunta.objects.filter(usuario__user=user).count()
        preguntas_por_usuario[user.username] = {
            'cantidad': preguntas_count,
            'is_active': user.userprofile.is_active  # Acceder al estado del perfil
        }
    
    context = {
        'universidades_count': Universidad.objects.count(),
        'cursos_count': Curso.objects.count(),
        'temas_count': Tema.objects.count(),
        'preguntas_count': Pregunta.objects.count(),
        'preguntas_por_usuario': preguntas_por_usuario,
        'preguntas_recientes': preguntas_recientes,
        'temas': Tema.objects.all(),
        'universidades': Universidad.objects.all(),
        'cursos': Curso.objects.all(),
    }
    
    # Agregar los valores de filtro al contexto
    for campo in ['tema', 'universidad', 'curso']:
        context[f'{campo}_filter'] = request.GET.get(campo)
    
    return render(request, 'Preguntas/admin_dashboard.html', context)

@login_required
@staff_member_required
def toggle_user_status(request, username):
    user = get_object_or_404(User, username=username)
    user_profile = get_object_or_404(UserProfile, user=user)

    # Cambiar el estado de is_active
    user_profile.is_active = not user_profile.is_active
    user_profile.save()

    # Mensaje de éxito
    if user_profile.is_active:
        messages.success(request, f'La cuenta de {user.username} ha sido activada.')
    else:
        messages.warning(request, f'La cuenta de {user.username} ha sido desactivada.')

    # Redirigir de vuelta al dashboard
    return redirect('admin-dashboard')  # Asegúrate de que este nombre coincida con tu URL

@staff_member_required
def generar_examen(request):
    # Filtros
    filtros = {}
    for campo in ['tema', 'universidad', 'curso']:
        valor = request.GET.get(campo)
        if valor:
            filtros[campo + '__id'] = valor

    # Preguntas filtradas
    preguntas = Pregunta.objects.filter(**filtros).order_by('-fecha_creacion')
    carrito = request.session.get('carrito', [])

    if request.method == 'POST':
        is_ajax = request.headers.get('X-Requested-With') == 'XMLHttpRequest'
        
        # Obtener todos los IDs de preguntas
        pregunta_ids = request.POST.getlist('preguntas')  # Cambiado a getlist para obtener múltiples IDs

        # Añadir al carrito
        if 'add_to_cart' in request.POST:
            for pregunta_id in pregunta_ids:
                if pregunta_id and pregunta_id not in carrito:
                    carrito.append(pregunta_id)
            request.session['carrito'] = carrito

            if is_ajax:
                return JsonResponse({'success': True})
            else:
                return redirect('generar_examen')

        # Añadir pregunta desde la vista previa
        elif 'add_preview' in request.POST:
            pregunta_id = request.POST.get('pregunta_id')  # ID de la pregunta desde la vista previa
            if pregunta_id and pregunta_id not in carrito:
                carrito.append(pregunta_id)
            request.session['carrito'] = carrito

            if is_ajax:
                return JsonResponse({'success': True})
            else:
                return redirect('generar_examen')

        # Eliminar del carrito
        elif 'remove_from_cart' in request.POST:
            for pregunta_id in pregunta_ids:
                if pregunta_id in carrito:
                    carrito.remove(pregunta_id)
            request.session['carrito'] = carrito

            if is_ajax:
                return JsonResponse({'success': True})
            else:
                return redirect('generar_examen')

        # Vaciar el carrito
        elif 'vaciar_carrito' in request.POST:
            carrito.clear()  # Vaciar el carrito
            request.session['carrito'] = carrito

            if is_ajax:
                return JsonResponse({'success': True})
            else:
                return redirect('generar_examen')

        # Descargar preguntas del carrito
        elif 'download' in request.POST:
            preguntas_seleccionadas = Pregunta.objects.filter(id__in=carrito)
            if not preguntas_seleccionadas:
                messages.error(request, "No hay preguntas en el carrito para descargar.")
                return redirect('generar_examen')

            buffer = combinar_documentos(preguntas_seleccionadas)
            response = HttpResponse(
                buffer,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = 'attachment; filename="examen_generado.docx"'
            return response

    # Obtener preguntas en el carrito
    carrito_preguntas = Pregunta.objects.filter(id__in=carrito)

    # Mostrar formulario
    context = {
        'preguntas': preguntas,
        'temas': Tema.objects.all(),
        'universidades': Universidad.objects.all(),
        'cursos': Curso.objects.all(),
        'carrito': carrito_preguntas,
    }

    # Agregar los valores de filtro al contexto
    for campo in ['tema', 'universidad', 'curso']:
        context[f'{campo}_filter'] = request.GET.get(campo)

    return render(request, 'Preguntas/generar_examen.html', context)

@login_required
def export_preguntas_recientes(request):

    preguntas = Pregunta.objects.select_related('usuario__user', 'universidad', 'curso', 'tema').order_by('-fecha_creacion')

    response = HttpResponse(content_type='text/csv; charset=utf-8')
    response['Content-Disposition'] = 'attachment; filename="preguntas_recientes.csv"'

    # Agregar BOM para UTF-8 para compatibilidad con Excel
    response.write(u'﻿'.encode('utf8'))

    fieldnames = ['Usuario', 'Universidad', 'Curso', 'Tema', 'Nivel', 'Fecha de Creación']
    writer = csv.DictWriter(response, fieldnames=fieldnames, delimiter=",", quoting=csv.QUOTE_MINIMAL)

    writer.writeheader()
    writer.writerows({
        'Usuario': pregunta.usuario.user.username if pregunta.usuario else 'Desconocido',
        'Universidad': pregunta.universidad.nombre,
        'Curso': pregunta.curso.nombre,
        'Tema': pregunta.tema.nombre,
        'Nivel': pregunta.nivel,
        'Fecha de Creación': pregunta.fecha_creacion.strftime("%Y-%m-%d %H:%M:%S"),
    } for pregunta in preguntas)

    return response